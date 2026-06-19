"""
Add charts and graphs to ROI-Analysis.docx.

Creates matplotlib charts from the data in the ROI Analysis, saves them as PNGs,
then builds a new Word document from the markdown source with charts inserted
at appropriate locations.

The existing .docx is OLE2 format and cannot be opened by python-docx.
Instead we use pandoc to convert the .md to .docx, then insert charts.
"""

import os
import subprocess
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CHARTS_DIR = os.path.join(SCRIPT_DIR, 'charts_temp')
INPUT_MD = os.path.join(SCRIPT_DIR, 'ROI-Analysis.md')
INTERMEDIATE_DOCX = os.path.join(SCRIPT_DIR, 'ROI-Analysis-base.docx')
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, 'ROI-Analysis.docx')

# Check for reference.docx for table borders
REFERENCE_DOCX = os.path.join(SCRIPT_DIR, 'reference.docx')

os.makedirs(CHARTS_DIR, exist_ok=True)

# ---------- color palette (Microsoft-inspired) ----------
BLUE = '#0078D4'
DARK_BLUE = '#004578'
TEAL = '#008575'
GREEN = '#107C10'
ORANGE = '#D83B01'
RED = '#A80000'
PURPLE = '#5C2D91'
YELLOW = '#FFB900'
GRAY = '#737373'
LIGHT_GRAY = '#E1E1E1'

THEME_COLORS = [BLUE, ORANGE, TEAL, GREEN, PURPLE, RED, YELLOW, DARK_BLUE]

def style_chart(ax, title, xlabel=None, ylabel=None):
    """Apply consistent styling to a chart."""
    ax.set_title(title, fontsize=13, fontweight='bold', pad=12, color='#1a1a1a')
    if xlabel:
        ax.set_xlabel(xlabel, fontsize=10, color='#444')
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=10, color='#444')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#cccccc')
    ax.spines['bottom'].set_color('#cccccc')
    ax.tick_params(colors='#555', labelsize=9)
    ax.grid(axis='y', alpha=0.3, linestyle='--')


# ============================================================
# Chart 1: Volume Trends Line Chart
# ============================================================
def create_volume_trends_chart():
    themes_declining = {
        'Auth/Sign-in': [24, 15, 3],
        'Account Login': [19, 17, 10],
        'Cloud Access': [34, 26, 23],
        'Parameter Override': [5, 1, 0],
    }
    themes_increasing = {
        'Device Enrollment': [17, 45, 27],
        'Cert Rotation': [11, 31, 25],
    }

    periods = ['Mar–Apr', 'Apr–May', 'May–Jun']

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.5), gridspec_kw={'width_ratios': [3, 2]})

    # Left panel: declining themes
    colors_dec = [BLUE, TEAL, GREEN, PURPLE]
    for i, (theme, vals) in enumerate(themes_declining.items()):
        ax1.plot(periods, vals, marker='o', linewidth=2.5, markersize=8,
                 color=colors_dec[i], label=theme, zorder=3)
        # annotate endpoints
        ax1.annotate(f'{vals[-1]}', (periods[-1], vals[-1]),
                     textcoords="offset points", xytext=(8, 0),
                     fontsize=8, color=colors_dec[i], fontweight='bold')

    style_chart(ax1, 'Themes with Published Docs — Volume Declining',
                ylabel='Incident Count')
    ax1.legend(fontsize=8, loc='upper right', framealpha=0.9)
    ax1.set_ylim(bottom=-1)

    # Right panel: increasing themes (external events)
    colors_inc = [ORANGE, RED]
    for i, (theme, vals) in enumerate(themes_increasing.items()):
        ax2.plot(periods, vals, marker='s', linewidth=2.5, markersize=8,
                 color=colors_inc[i], label=theme, linestyle='--', zorder=3)
        ax2.annotate(f'{vals[-1]}', (periods[-1], vals[-1]),
                     textcoords="offset points", xytext=(8, 0),
                     fontsize=8, color=colors_inc[i], fontweight='bold')

    style_chart(ax2, 'Themes with External Events — Volume Increased',
                ylabel='Incident Count')
    ax2.legend(fontsize=8, loc='upper right', framealpha=0.9)

    fig.tight_layout(pad=2)
    path = os.path.join(CHARTS_DIR, 'volume_trends.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Chart 2: Signal vs. Noise Donut
# ============================================================
def create_signal_noise_chart():
    fig, ax = plt.subplots(figsize=(4, 4))
    sizes = [240, 115]  # 355 total, 240 signal
    labels = ['Signal Incidents\n(240)', 'Noise / Automated\n(115)']
    colors = [BLUE, LIGHT_GRAY]
    explode = (0.03, 0)

    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=colors, explode=explode,
        autopct='%1.0f%%', startangle=90, pctdistance=0.75,
        textprops={'fontsize': 10})
    for at in autotexts:
        at.set_fontweight('bold')
        at.set_fontsize(11)

    # donut hole
    centre = plt.Circle((0, 0), 0.50, fc='white')
    ax.add_artist(centre)
    ax.text(0, 0, '355\ntotal', ha='center', va='center',
            fontsize=14, fontweight='bold', color='#333')

    ax.set_title('Incident Signal vs. Noise (May–Jun)', fontsize=12,
                 fontweight='bold', pad=15, color='#1a1a1a')

    path = os.path.join(CHARTS_DIR, 'signal_noise.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Chart 3: Severity Distribution Donut
# ============================================================
def create_severity_chart():
    # 355 total: 88% Sev3, 1.4% Sev2, 0.3% Sev1, remainder Sev4
    sev3 = round(355 * 0.88)   # 312
    sev2 = round(355 * 0.014)  # 5
    sev1 = round(355 * 0.003)  # 1
    sev4 = 355 - sev3 - sev2 - sev1  # 37

    fig, ax = plt.subplots(figsize=(6, 3.5))

    severities = ['Sev1', 'Sev2', 'Sev3', 'Sev4']
    counts = [sev1, sev2, sev3, sev4]
    colors = [RED, ORANGE, BLUE, GRAY]
    pcts = [f'{c/355*100:.1f}%' for c in counts]

    bars = ax.barh(severities, counts, color=colors, height=0.55, zorder=3)

    # Annotate each bar with count and percentage
    for bar, count, pct in zip(bars, counts, pcts):
        # For bars that are too short for internal labels, put label outside
        if count < 50:
            ax.text(bar.get_width() + 5, bar.get_y() + bar.get_height()/2,
                    f'{count}  ({pct})', va='center', fontsize=9, color='#444')
        else:
            ax.text(bar.get_width() - 5, bar.get_y() + bar.get_height()/2,
                    f'{count}  ({pct})', va='center', ha='right', fontsize=9,
                    color='white', fontweight='bold')

    style_chart(ax, 'Severity Distribution (May-Jun)',
                xlabel='Number of Incidents')
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.grid(axis='y', visible=False)
    ax.invert_yaxis()

    # Add callout
    ax.text(sev3 * 0.55, 3.7, 'Sev3 drives ~89% of total weighted benefit potential',
            fontsize=8.5, fontstyle='italic', color=BLUE)

    ax.set_xlim(right=max(counts) * 1.15)
    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, 'severity_dist.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Chart 4: Time-to-Mitigate Horizontal Bar Chart
# ============================================================
def create_ttm_chart():
    themes = ['Cert Rotation', 'CLI/Client', 'Auth Issues',
              'Subscription Mgmt', 'Device Enrollment', 'Account Login',
              'Cloud Access', 'Approval Routing']
    ttm = [6.2, 7.7, 20.9, 21.9, 23.1, 74.1, 83.7, 114.6]
    counts = [7, 5, 18, 13, 24, 8, 38, 44]
    has_docs = [True, True, True, False, True, True, True, True]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    colors = [BLUE if d else GRAY for d in has_docs]
    bars = ax.barh(themes, ttm, color=colors, height=0.6, zorder=3)

    # annotate with count and hours
    for bar, t, n in zip(bars, ttm, counts):
        ax.text(bar.get_width() + 2, bar.get_y() + bar.get_height()/2,
                f'{t}h  (n={n})', va='center', fontsize=8.5, color='#444')

    style_chart(ax, 'Median Time-to-Mitigate by Theme (May–Jun)',
                xlabel='Median Hours to Mitigate')
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.grid(axis='y', visible=False)

    # legend for docs status
    from matplotlib.patches import Patch
    legend_elements = [Patch(facecolor=BLUE, label='Docs published'),
                       Patch(facecolor=GRAY, label='No docs yet')]
    ax.legend(handles=legend_elements, fontsize=8, loc='lower right')

    ax.set_xlim(right=max(ttm) * 1.25)
    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, 'ttm_by_theme.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Chart 5: Preventability Horizontal Bar Chart
# ============================================================
def create_preventability_chart():
    fig, ax = plt.subplots(figsize=(7, 3.5))

    categories = [
        'Doc helped but\ncode fix moot',
        'Unclear\n(insufficient detail)',
        'Doc would have\nhelped (gap)',
        'Docs exist, user\ndidn\'t search',
        'Engineering fix\n(docs irrelevant)',
    ]
    counts = [1, 2, 2, 4, 9]  # sorted ascending for bottom-to-top
    colors = [YELLOW, LIGHT_GRAY, GREEN, ORANGE, GRAY]
    total = sum(counts)

    bars = ax.barh(categories, counts, color=colors, height=0.6, zorder=3,
                   edgecolor='white', linewidth=1)

    for bar, count in zip(bars, counts):
        pct = count / total * 100
        ax.text(bar.get_width() + 0.15, bar.get_y() + bar.get_height()/2,
                f'{count}  ({pct:.0f}%)', va='center', fontsize=9, color='#444')

    style_chart(ax, 'Human-Calibrated Preventability Assessment (n=18)',
                xlabel='Number of Incidents')
    ax.set_xlim(0, max(counts) * 1.4)
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.grid(axis='y', visible=False)

    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, 'preventability.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Chart 6: Token Cost vs. Incident Value
# ============================================================
def create_token_cost_chart():
    fig, ax = plt.subplots(figsize=(7, 4.5))

    price_labels = ['$3\n(current)', '$10\n(post-subsidy)', '$25\n(high)', '$50\n(scarce)', '$100\n(extreme)']
    token_costs = [7.68, 25.61, 64.03, 128.05, 256.10]

    # Value of prevented incidents: 6 incidents x 3 hrs x $180/hr fully loaded
    monthly_value = 6 * 3 * 180  # $3,240
    single_icm = 3 * 180         # $540

    x = np.arange(len(price_labels))
    width = 0.5

    bars = ax.bar(x, token_costs, width, color=BLUE, zorder=3, label='Monthly token cost')

    # value line
    ax.axhline(monthly_value, color=GREEN, linestyle='--', alpha=0.6, linewidth=2, zorder=2)
    ax.text(len(x) - 0.5, monthly_value + 80,
            f'Value of ~6 prevented incidents (${monthly_value:,}/mo)',
            fontsize=8.5, color=GREEN, fontweight='bold', ha='right')

    # single incident line
    ax.axhline(single_icm, color=ORANGE, linestyle=':', alpha=0.7, linewidth=1.5, zorder=2)
    ax.text(len(x) - 0.5, single_icm + 50,
            f'Value of 1 prevented incident (${single_icm})',
            fontsize=8, color=ORANGE, ha='right')

    # annotate bars
    for i, v in enumerate(token_costs):
        ax.text(i, v + 15, f'${v:.0f}', ha='center', fontsize=8, fontweight='bold', color=BLUE)

    style_chart(ax, 'Monthly Token Cost vs. Value of Prevented Incidents',
                xlabel='Assumed $/Million Tokens', ylabel='Monthly Value ($)')
    ax.set_xticks(x)
    ax.set_xticklabels(price_labels)
    ax.set_ylim(0, monthly_value * 1.15)
    ax.grid(axis='y', alpha=0.2)

    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, 'token_cost_vs_value.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Chart 7: Break-Even Sensitivity
# ============================================================
def create_breakeven_chart():
    fig, ax = plt.subplots(figsize=(6, 4))

    oce_time_assumptions = [1, 2, 3, 4, 5]
    # Break-even = operating cost (3.5 hrs) / engineer hours per incident
    # At 3.5 hrs operating cost and X hrs/incident: need 3.5/X incidents
    icms_needed = [3.5 / t for t in oce_time_assumptions]

    bars = ax.bar(range(len(oce_time_assumptions)), icms_needed,
                  color=BLUE, width=0.5, zorder=3)

    # Annotate bars
    for i, (bar, val) in enumerate(zip(bars, icms_needed)):
        ax.text(i, val + 0.08, f'{val:.1f}', ha='center', fontsize=9,
                fontweight='bold', color=BLUE)

    # Estimated prevented incidents line
    ax.axhline(6, color=GREEN, linestyle='--', linewidth=2, alpha=0.7, zorder=2)
    ax.text(len(oce_time_assumptions) - 0.6, 6.2,
            'Estimated ~6 incidents\nprevented/month',
            fontsize=8.5, color=GREEN, fontweight='bold', ha='right')

    style_chart(ax, 'Break-Even Sensitivity: Incidents Needed by Engineer Time Assumption',
                xlabel='Assumed engineer hours per incident', ylabel='Incidents/month needed to break even')
    ax.set_xticks(range(len(oce_time_assumptions)))
    ax.set_xticklabels([f'{t}h' for t in oce_time_assumptions])
    ax.set_ylim(0, 7.5)
    ax.grid(axis='y', alpha=0.2)

    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, 'breakeven_sensitivity.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Chart 8: Development Cost Payback
# ============================================================
def create_dev_payback_chart():
    fig, ax = plt.subplots(figsize=(7, 4.5))

    price_labels = ['$3\n(current)', '$10\n(post-subsidy)', '$25\n(high)', '$50\n(scarce)', '$100\n(extreme)']
    dev_costs = [780, 2600, 6500, 13000, 26000]  # 260M tokens at each price level

    # Monthly net value: ~6 prevented incidents * ~3hrs * $180/hr fully loaded = ~$3,240/mo
    # Minus monthly token cost at each price level
    monthly_token_costs = [7.68, 25.61, 64.03, 128.05, 256.10]
    monthly_value = 6 * 3 * 180  # $3,240
    monthly_net = [monthly_value - tc for tc in monthly_token_costs]

    # Payback months = dev cost / monthly net value
    payback_months = [dc / mn for dc, mn in zip(dev_costs, monthly_net)]

    x = np.arange(len(price_labels))
    bars = ax.bar(x, payback_months, width=0.5, color=PURPLE, zorder=3)

    # Annotate bars
    for i, (bar, months) in enumerate(zip(bars, payback_months)):
        label = f'{months:.1f} mo' if months < 12 else f'{months:.0f} mo\n({months/12:.1f} yr)'
        ax.text(i, bar.get_height() + 0.3, label, ha='center', fontsize=8.5,
                fontweight='bold', color=PURPLE)

    style_chart(ax, 'Development Cost Payback Period',
                xlabel='Assumed $/Million Tokens',
                ylabel='Months to Recover Dev Token Cost')
    ax.set_xticks(x)
    ax.set_xticklabels(price_labels)
    ax.set_ylim(0, max(payback_months) * 1.2)
    ax.grid(axis='y', alpha=0.2)

    # Add context line at 12 months
    if max(payback_months) > 12:
        ax.axhline(12, color=ORANGE, linestyle=':', linewidth=1.5, alpha=0.7)
        ax.text(0.1, 12.5, '1 year', fontsize=8, color=ORANGE)

    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, 'dev_payback.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Insert charts into the Word document
# ============================================================
def find_paragraph_containing(doc, text):
    """Find the index of a paragraph containing the given text."""
    # Normalize dashes for matching
    import re
    normalized_text = text.replace('\u2013', '-').replace('\u2014', '-')
    for i, p in enumerate(doc.paragraphs):
        normalized_para = p.text.replace('\u2013', '-').replace('\u2014', '-')
        if normalized_text in normalized_para:
            return i
    return None


def add_table_borders(doc):
    """Add borders to all tables in the document."""
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        
        borders_xml = f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        </w:tblBorders>'''
        
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        
        tblPr.append(parse_xml(borders_xml))
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)


def insert_chart_after_paragraph(doc, para_idx, image_path, caption, width_inches=5.5):
    """Insert a chart image and caption after a given paragraph index."""
    # We insert after para_idx by manipulating the XML directly
    ref_para = doc.paragraphs[para_idx]

    # Create image paragraph
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    # Create caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    cap_run.font.italic = True

    # Move the new paragraphs to after ref_para
    # In python-docx, newly added paragraphs go to the end.
    # We need to move them in the XML tree.
    ref_element = ref_para._element
    parent = ref_element.getparent()

    # Move image paragraph
    img_element = new_para._element
    parent.remove(img_element)
    ref_element.addnext(img_element)

    # Move caption after image
    cap_element = cap_para._element
    parent.remove(cap_element)
    img_element.addnext(cap_element)


def insert_charts_into_docx(chart_paths):
    """Convert MD to docx via pandoc, then insert charts at appropriate locations."""
    # Step 1: Convert markdown to docx via pandoc
    pandoc_cmd = ['pandoc', INPUT_MD, '-o', INTERMEDIATE_DOCX, '--from=markdown', '--to=docx']
    # Only use reference.docx if it's a valid OOXML file (not OLE2)
    if os.path.exists(REFERENCE_DOCX):
        with open(REFERENCE_DOCX, 'rb') as f:
            magic = f.read(2)
        if magic == b'PK':
            pandoc_cmd.extend(['--reference-doc', REFERENCE_DOCX])
            print(f"  Using reference.docx for table borders")
        else:
            print(f"  reference.docx is OLE2 format, skipping (will add borders programmatically)")
    
    result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  Pandoc error: {result.stderr}")
        raise RuntimeError("Pandoc conversion failed")
    print(f"  Converted {INPUT_MD} to docx via pandoc")

    doc = Document(INTERMEDIATE_DOCX)
    
    # Add borders to all tables
    add_table_borders(doc)
    print("  Added borders to all tables")

    # Define insertion points: (search_text, chart_key, caption, width)
    # Ordered by document position (TTM first at ~para 41, then volume trends ~47, etc.)
    insertions = [
        # Figure 1: TTM - after the TTM table (appears first in doc)
        ('Approval Routing and Cloud Access represent the highest time-burden',
         'ttm_by_theme',
         'Figure 1: Median time-to-mitigate by theme. Blue = docs published; gray = no docs yet.',
         5.5),

        # Figure 2: Volume trends - after the interpretation paragraph
        ('These trends are consistent with the hypothesis',
         'volume_trends',
         'Figure 2: Incident volume trends across three measurement periods. Themes with published docs generally declined, while increases are concentrated in themes affected by known external events. Trends are consistent with, but do not prove, documentation impact.',
         6.0),

        # Figure 3: Preventability - after key findings
        ('50% required engineering fixes regardless',
         'preventability',
         'Figure 3: Human-calibrated preventability (n=18). Most incidents required engineering fixes, but a small doc-preventable share is sufficient to clear the low break-even threshold.',
         5.5),

        # Figure 4: Break-even sensitivity - after the sensitivity table
        ('The break-even threshold is low enough',
         'breakeven_sensitivity',
         'Figure 4: Break-even requires fewer than 2 incidents/month under baseline assumptions (3 hrs/incident). Green line shows estimated actual prevention rate (~6/month).',
         5.0),

        # Figure 5: Token cost - after the token pricing table
        ('Token cost does not become the binding constraint',
         'token_cost_vs_value',
         'Figure 5: Even at extreme token prices, monthly operating cost remains far below the value of prevented incidents. Assumes ~6 prevented incidents/month, 3 engineer hrs/incident, $180/hr fully loaded cost.',
         5.5),
    ]

    # Process insertions in reverse document order so indices don't shift
    insertion_data = []
    for search_text, chart_key, caption, width in insertions:
        idx = find_paragraph_containing(doc, search_text)
        if idx is not None:
            insertion_data.append((idx, chart_paths[chart_key], caption, width))
            print(f"  Found '{search_text[:50]}...' at paragraph {idx}")
        else:
            print(f"  WARNING: Could not find '{search_text[:50]}...'")

    # Sort by index descending so we insert from bottom to top
    insertion_data.sort(key=lambda x: x[0], reverse=True)

    for idx, img_path, caption, width in insertion_data:
        insert_chart_after_paragraph(doc, idx, img_path, caption, width)
        print(f"  Inserted chart after paragraph {idx}")

    doc.save(OUTPUT_DOCX)
    print(f"\nSaved: {OUTPUT_DOCX}")


# ============================================================
# Main
# ============================================================
if __name__ == '__main__':
    print("Generating charts...")

    chart_paths = {}
    chart_paths['volume_trends'] = create_volume_trends_chart()
    print("  [ok] Volume trends line chart")

    chart_paths['ttm_by_theme'] = create_ttm_chart()
    print("  [ok] Time-to-mitigate bar chart")

    chart_paths['preventability'] = create_preventability_chart()
    print("  [ok] Preventability stacked bar")

    chart_paths['breakeven_sensitivity'] = create_breakeven_chart()
    print("  [ok] Break-even sensitivity chart")

    chart_paths['token_cost_vs_value'] = create_token_cost_chart()
    print("  [ok] Token cost vs. value chart")

    print(f"\nAll charts saved to {CHARTS_DIR}")
    print("\nInserting charts into Word document...")
    insert_charts_into_docx(chart_paths)

    # Clean up temp files
    import shutil
    shutil.rmtree(CHARTS_DIR, ignore_errors=True)
    if os.path.exists(INTERMEDIATE_DOCX):
        os.remove(INTERMEDIATE_DOCX)
    print("Cleaned up temporary files.")
    print("Done!")
