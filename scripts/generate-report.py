"""
Incident Doc Gap Report Generator
==================================

Generates a Word (.docx) Doc Gap Report from the JSON artifacts produced
by Pass 1 (extraction) and Pass 2 (classification) of the Incident Doc Gap
Analysis workflow.

Usage:
    python generate-report.py --period "May-Jun 2026" --start 2026-05-12 --end 2026-06-01

Inputs (read from --data-dir):
    - Incident-Extract-{period}.jsonl     (Pass 1 output)
    - priority-matrix.json                (Pass 2 output)
    - gap-classification.json             (Pass 2 output)
    - theme-summary.json                  (Pass 2 output)

Output:
    - Doc-Gap-Report-{period}.docx        (in --output-dir)

Requires: python-docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt

# Configurable: set this to your incident portal base URL
INCIDENT_BASE_URL = "{INCIDENT_PORTAL_URL}"

# Avoidable rate formulas (fixed per workflow spec)
AVOIDABLE_RATES = {
    "real_doc_gap": 0.40,
    "partial_doc_gap": 0.30,
    "discoverability_problem": 0.20,
    "engineering_fix": 0.0,
    "operational_gap": 0.0,
    "no_doc_action_needed": 0.0,
}

CLASSIFICATION_LABELS = {
    "real_doc_gap": "Real doc gap",
    "partial_doc_gap": "Partial doc gap",
    "discoverability_problem": "Discoverability problem",
    "engineering_fix": "Engineering fix",
    "operational_gap": "Operational gap",
    "no_doc_action_needed": "No doc action needed",
}

FIX_EFFORT_MAP = {
    "real_doc_gap": "High",
    "partial_doc_gap": "Medium",
    "discoverability_problem": "Low (discoverability)",
    "engineering_fix": "N/A (engineering)",
    "operational_gap": "N/A (operational)",
    "no_doc_action_needed": "N/A",
}


# ─── Utility: JSON / JSONL ─────────────────────────────────────────────────────

def load_json(path: Path):
    if not path.exists():
        print(f"ERROR: Required file not found: {path}", file=sys.stderr)
        sys.exit(1)
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in {path}: {e}", file=sys.stderr)
        sys.exit(1)


def load_jsonl(path: Path) -> list[dict]:
    if not path.exists():
        print(f"ERROR: Required file not found: {path}", file=sys.stderr)
        sys.exit(1)
    rows = []
    with path.open(encoding="utf-8") as f:
        for lineno, line in enumerate(f, 1):
            line = line.strip()
            if line:
                try:
                    rows.append(json.loads(line))
                except json.JSONDecodeError as e:
                    print(f"ERROR: Invalid JSON on line {lineno} of {path}: {e}", file=sys.stderr)
                    sys.exit(1)
    return rows


def get_motivated_incidents_for_theme(slug: str, draft_dir: Path) -> list[str]:
    """Extract motivated_by_icms IDs from a draft's YAML frontmatter."""
    if not draft_dir.exists():
        return []
    md_file = None
    for f in draft_dir.glob("*.md"):
        if slug.replace("-", "") in f.stem.replace("-", "") \
                and "discoverability" not in f.stem:
            md_file = f
            break
    if not md_file:
        return []
    text = md_file.read_text(encoding="utf-8")
    ids = []
    for match in re.finditer(r'- id:\s*(\d+)', text):
        ids.append(match.group(1))
    return ids


# ─── Utility: Word formatting ──────────────────────────────────────────────────

def set_font(style, name: str, size: int | None = None, bold: bool | None = None):
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if style.element.rPr is not None:
        r_fonts = style.element.rPr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            style.element.rPr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)


def configure_styles(document: Document) -> None:
    set_font(document.styles["Normal"], "Aptos", 10)
    set_font(document.styles["Title"], "Aptos", 20, True)
    set_font(document.styles["Heading 1"], "Aptos", 15, True)
    set_font(document.styles["Heading 2"], "Aptos", 13, True)
    set_font(document.styles["Heading 3"], "Aptos", 11, True)


def add_shading(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_table_headers(table) -> None:
    for cell in table.rows[0].cells:
        add_shading(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _new_hyperlink_run(text: str):
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    return run


def add_external_hyperlink(paragraph, text: str, url: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True))
    hyperlink.append(_new_hyperlink_run(text))
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    hyperlink.append(_new_hyperlink_run(text))
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9.5)


def incident_url(incident_id: str) -> str:
    return f"{INCIDENT_BASE_URL}{incident_id}"


def add_incident_links(paragraph, incident_ids: Iterable[str]) -> None:
    ids = list(incident_ids)
    for index, inc_id in enumerate(ids):
        add_external_hyperlink(paragraph, str(inc_id), incident_url(str(inc_id)))
        if index < len(ids) - 1:
            paragraph.add_run(", ")


def add_labelled_text(document: Document, label: str, text: str):
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(text)
    return paragraph


# ─── Priority computation ──────────────────────────────────────────────────────

def compute_priority(classification: str, count: int) -> str:
    if classification == "real_doc_gap" and count >= 5:
        return "P1"
    if classification == "partial_doc_gap" and count >= 10:
        return "P1"
    if classification == "partial_doc_gap":
        return "P2"
    if classification == "discoverability_problem" and count >= 15:
        return "P2"
    if classification == "discoverability_problem":
        return "P3"
    return "P4"


def build_priority_rows(theme_summary: dict, gap_classification: dict) -> list[dict]:
    rows = []
    for slug, data in theme_summary.items():
        if slug == "unclassified":
            continue
        gap = gap_classification.get(slug)
        if not gap:
            continue

        signal_count = data.get("signal_count", data["count"])
        count = data["count"]
        cls = gap["classification"]
        if cls not in AVOIDABLE_RATES:
            raise ValueError(f"Unknown classification '{cls}' for theme '{slug}'.")
        rate = AVOIDABLE_RATES[cls]
        avoidable = int(round(signal_count * rate + 0.4999))
        priority = compute_priority(cls, signal_count)

        rows.append({
            "slug": slug,
            "priority": priority,
            "theme": data["display_name"],
            "incidents": count,
            "signal_count": signal_count,
            "avoidable": f"{avoidable} (~{int(rate*100)}%)",
            "avoidable_count": avoidable,
            "fix_effort": FIX_EFFORT_MAP.get(cls, "Unknown"),
            "classification": cls,
            "needs_draft": cls in ("real_doc_gap", "partial_doc_gap"),
        })

    rows.sort(key=lambda r: (r["priority"], -r["incidents"]))
    return rows


# ─── Report sections ──────────────────────────────────────────────────────────

def add_executive_summary(document: Document, period: str, start: str, end: str,
                          raw_incidents: list[dict], rows: list[dict],
                          theme_summary: dict) -> None:
    document.add_paragraph("Executive Summary", style="Heading 1")

    signal = sum(1 for r in raw_incidents if r.get("SignalClass") == "signal")
    automated = sum(1 for r in raw_incidents if r.get("SignalClass") == "automated")
    noise = sum(1 for r in raw_incidents if r.get("SignalClass") == "noise")
    total = len(raw_incidents)

    team_counts = {}
    for r in raw_incidents:
        key = f"{r['TeamName']} ({r['TeamId']})"
        team_counts[key] = team_counts.get(key, 0) + 1

    total_avoidable = sum(r["avoidable_count"] for r in rows)

    points = [
        f"Date range: {start} to {end} (UTC)",
        f"Total incidents: {total} ({signal} signal, {automated} automated, {noise} noise)",
        f"Teams: {', '.join(f'{k}: {v}' for k, v in team_counts.items())}",
        f"Themes classified: {len(rows)}",
        f"Total estimated avoidable: {total_avoidable}",
    ]
    for item in points:
        document.add_paragraph(item, style="List Bullet")


def add_priority_matrix(document: Document, rows: list[dict], draft_dir: Path) -> None:
    document.add_paragraph("Priority Matrix", style="Heading 1")

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Priority", "Theme", "# Incidents", "Avoidable (est.)", "Fix Effort", "Draft Fix", "Status"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header)
    format_table_headers(table)

    for row in rows:
        tr = table.add_row().cells
        set_cell_text(tr[0], row["priority"])
        set_cell_text(tr[1], row["theme"])
        set_cell_text(tr[2], str(row["incidents"]))
        set_cell_text(tr[3], row["avoidable"])
        set_cell_text(tr[4], row["fix_effort"])
        tr[5].paragraphs[0].add_run("—" if not row["needs_draft"] else "See drafts")
        set_cell_text(tr[6], "—")


def add_metadata_table(document: Document, metadata: dict) -> None:
    document.add_paragraph("Reproducibility Metadata", style="Heading 1")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, ["Field", "Value"]):
        set_cell_text(cell, header)
    format_table_headers(table)

    for field, value in metadata.items():
        row = table.add_row().cells
        set_cell_text(row[0], field)
        set_cell_text(row[1], str(value))


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate Incident Doc Gap Report (.docx)")
    parser.add_argument("--period", required=True, help="Report period label, e.g. 'May-Jun 2026'")
    parser.add_argument("--start", required=True, help="Start date (UTC)")
    parser.add_argument("--end", required=True, help="End date (UTC)")
    parser.add_argument("--data-dir", required=True, help="Path to run data directory")
    parser.add_argument("--output-dir", required=True, help="Path to output directory")
    args = parser.parse_args()

    data_dir = Path(args.data_dir)
    output_dir = Path(args.output_dir)
    period = args.period
    drafts_dir = output_dir / "doc-fix-drafts"

    if not data_dir.exists():
        print(f"ERROR: Data directory not found: {data_dir}", file=sys.stderr)
        sys.exit(1)
    if not output_dir.exists():
        print(f"ERROR: Output directory not found: {output_dir}", file=sys.stderr)
        sys.exit(1)

    # Load data
    jsonl_path = output_dir / f"Incident-Extract-{period}.jsonl"
    if not jsonl_path.exists():
        print(f"ERROR: JSONL not found at {jsonl_path}", file=sys.stderr)
        sys.exit(1)

    raw_incidents = load_jsonl(jsonl_path)
    priority_matrix = load_json(data_dir / "priority-matrix.json")
    gap_classification = load_json(data_dir / "gap-classification.json")
    theme_summary = load_json(data_dir / "theme-summary.json")

    rows = build_priority_rows(theme_summary, gap_classification)

    metadata = {
        "JSONL source filename": jsonl_path.name,
        "Extraction timestamp": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "Date range (UTC)": f"{args.start}T00:00:00Z – {args.end}T00:00:00Z",
        "Total incidents (raw)": str(len(raw_incidents)),
        "Signal incidents": str(sum(1 for r in raw_incidents if r.get("SignalClass") == "signal")),
        "Theme taxonomy version": "See data/theme-taxonomy.yml",
    }

    # Generate document
    document = Document()
    configure_styles(document)

    title = document.add_paragraph(f"Incident Documentation Gap Report — {period}", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_executive_summary(document, period, args.start, args.end, raw_incidents, rows, theme_summary)
    add_priority_matrix(document, rows, drafts_dir)
    add_metadata_table(document, metadata)

    # Save
    output_path = output_dir / f"Doc-Gap-Report-{period}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"Saved report to {output_path}")


if __name__ == "__main__":
    main()
